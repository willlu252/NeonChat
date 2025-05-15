# backend/file_utils.py
import os
import base64
import io
from typing import Optional, Tuple
try:
    import docx
except ImportError:
    try:
        from python_docx import Document as docx
    except ImportError:
        # Create a mock docx module for basic functionality
        class MockDocument:
            def __init__(self, *args, **kwargs):
                self.paragraphs = []
        
        class MockDocx:
            def Document(self, *args, **kwargs):
                return MockDocument()
        
        docx = MockDocx()
        print("WARNING: python-docx module not found. DOCX processing will be limited.")
import re

def extract_text_from_docx(docx_content: str) -> str:
    """
    Extract text from a DOCX file content (base64 encoded).
    
    Args:
        docx_content: Base64 encoded DOCX file content
        
    Returns:
        Extracted text from the DOCX file
    """
    try:
        # Extract the base64 content (remove data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64, prefix)
        if isinstance(docx_content, str) and docx_content.startswith('data:'):
            # Extract the content type and base64 data
            content_parts = docx_content.split(';base64,')
            if len(content_parts) == 2:
                base64_data = content_parts[1]
                # Decode base64 data
                binary_data = base64.b64decode(base64_data)
                # Create a file-like object
                docx_file = io.BytesIO(binary_data)
                # Open the DOCX file
                doc = docx.Document(docx_file)
                # Extract text from paragraphs
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
                return text
        
        return "Error: Invalid DOCX content format"
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def process_file_content(file_content: str, file_type: str) -> Tuple[bool, str]:
    """
    Process file content based on file type.
    
    Args:
        file_content: Base64 encoded file content
        file_type: MIME type of the file
        
    Returns:
        Tuple of (success, text)
    """
    if file_type.startswith('application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
        # DOCX file
        text = extract_text_from_docx(file_content)
        return (True, text)
    elif file_type.startswith('text/'):
        # Text file
        try:
            if isinstance(file_content, str) and file_content.startswith('data:'):
                # Extract the content type and base64 data
                content_parts = file_content.split(';base64,')
                if len(content_parts) == 2:
                    base64_data = content_parts[1]
                    # Decode base64 data
                    binary_data = base64.b64decode(base64_data)
                    # Decode text
                    text = binary_data.decode('utf-8')
                    return (True, text)
        except Exception as e:
            return (False, f"Error extracting text from text file: {str(e)}")
    
    return (False, f"Unsupported file type: {file_type}")
